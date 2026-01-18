"""
文档解析服务
支持多种格式的文档解析：PDF、Excel、Word、网页、图片、TXT、数据库
"""
from __future__ import annotations

import os
import io
from typing import Optional, Dict, Any
from pathlib import Path

from ..utils import load_env

# PDF 解析
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    pdfplumber = None

try:
    import fitz  # PyMuPDF (备用)
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

# Word 文档解析
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None

# Excel 解析
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None

# 网页提取
try:
    import trafilatura
    TRAFILATURA_AVAILABLE = True
except ImportError:
    TRAFILATURA_AVAILABLE = False
    trafilatura = None

# 图片 OCR
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    PaddleOCR = None

# 备用 OCR (pytesseract)
try:
    from PIL import Image
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    pytesseract = None
    Image = None


def parse_pdf(file_data: bytes, filename: Optional[str] = None) -> str:
    """
    解析 PDF 文件（使用 pdfplumber，可保留表格）
    如果 pdfplumber 不可用，回退到 PyMuPDF
    """
    load_env()
    max_pages = int(os.environ.get("PDF_MAX_PAGES", "50"))
    max_chars = int(os.environ.get("PDF_MAX_CHARS", "50000"))
    
    text_parts = []
    
    # 优先使用 pdfplumber（可保留表格）
    if PDFPLUMBER_AVAILABLE:
        try:
            pdf_file = io.BytesIO(file_data)
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = min(len(pdf.pages), max_pages)
                
                for i in range(total_pages):
                    page = pdf.pages[i]
                    
                    # 提取文本
                    page_text = page.extract_text() or ""
                    if page_text:
                        text_parts.append(f"=== 第 {i+1} 页 ===\n{page_text}\n")
                    
                    # 提取表格
                    tables = page.extract_tables()
                    if tables:
                        for table_idx, table in enumerate(tables):
                            if table:
                                table_text = "表格内容：\n"
                                for row in table:
                                    if row:
                                        # 过滤 None 值
                                        row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                        table_text += row_text + "\n"
                                text_parts.append(f"=== 第 {i+1} 页 表格 {table_idx+1} ===\n{table_text}\n")
                    
                    # 检查字符限制
                    current_text = "".join(text_parts)
                    if len(current_text) >= max_chars:
                        break
                
                result = "".join(text_parts)
                if len(result) > max_chars:
                    result = result[:max_chars]
                return result.strip()
        except Exception as e:
            print(f"⚠ pdfplumber 解析失败: {e}")
    
    # 回退到 PyMuPDF
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
            n = min(getattr(doc, "page_count", len(doc)), max_pages)
            for i in range(n):
                try:
                    page = doc.load_page(i)
                    t = page.get_text() or ""
                    if t:
                        text_parts.append(f"=== 第 {i+1} 页 ===\n{t}\n")
                    if len("".join(text_parts)) >= max_chars:
                        break
                except Exception:
                    continue
            try:
                doc.close()
            except Exception:
                pass
            
            result = "".join(text_parts)
            if len(result) > max_chars:
                result = result[:max_chars]
            return result.strip()
        except Exception as e:
            print(f"⚠ PyMuPDF 解析失败: {e}")
    
    return ""


def parse_word(file_data: bytes, filename: Optional[str] = None) -> str:
    """解析 Word 文档 (.docx)"""
    if not PYTHON_DOCX_AVAILABLE:
        return ""
    
    try:
        doc_file = io.BytesIO(file_data)
        doc = Document(doc_file)
        
        text_parts = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text + "\n")
        
        # 提取表格
        for table in doc.tables:
            table_text = "表格内容：\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                table_text += row_text + "\n"
            text_parts.append(table_text + "\n")
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"⚠ Word 文档解析失败: {e}")
        return ""


def parse_excel(file_data: bytes, filename: Optional[str] = None) -> str:
    """解析 Excel 文档 (.xlsx, .xls)"""
    if not PANDAS_AVAILABLE:
        return ""
    
    try:
        excel_file = io.BytesIO(file_data)
        
        # 尝试读取所有工作表
        excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
        
        text_parts = []
        
        for sheet_name, df in excel_data.items():
            text_parts.append(f"=== 工作表: {sheet_name} ===\n")
            
            # 转换为文本格式（保留表格结构）
            # 使用制表符分隔，便于阅读
            text_parts.append(df.to_string(index=False))
            text_parts.append("\n\n")
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"⚠ Excel 文档解析失败: {e}")
        # 尝试使用 xlrd 引擎（.xls 格式）
        try:
            excel_file = io.BytesIO(file_data)
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
            text_parts = []
            for sheet_name, df in excel_data.items():
                text_parts.append(f"=== 工作表: {sheet_name} ===\n")
                text_parts.append(df.to_string(index=False))
                text_parts.append("\n\n")
            return "\n".join(text_parts).strip()
        except Exception as e2:
            print(f"⚠ Excel 文档解析失败（xlrd）: {e2}")
            return ""


def parse_webpage(url: str) -> str:
    """解析网页内容（使用 trafilatura）"""
    if not TRAFILATURA_AVAILABLE:
        return ""
    
    try:
        # 下载并提取网页内容
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            text = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
            return text or ""
        return ""
    except Exception as e:
        print(f"⚠ 网页解析失败: {e}")
        return ""


def parse_image(file_data: bytes, filename: Optional[str] = None) -> str:
    """
    解析图片文件（使用 PaddleOCR，回退到 pytesseract）
    """
    load_env()
    
    # 优先使用 PaddleOCR
    if PADDLEOCR_AVAILABLE:
        try:
            ocr = PaddleOCR(use_angle_cls=True, lang='ch')
            result = ocr.ocr(file_data, cls=True)
            
            text_parts = []
            if result and result[0]:
                for line in result[0]:
                    if line and len(line) >= 2:
                        text_info = line[1]
                        if text_info and len(text_info) >= 2:
                            text_parts.append(text_info[0])
            
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"⚠ PaddleOCR 解析失败: {e}")
    
    # 回退到 pytesseract
    if PYTESSERACT_AVAILABLE and Image:
        try:
            img = Image.open(io.BytesIO(file_data))
            
            # 设置 Tesseract 命令路径（如果配置了）
            try:
                cmd = os.environ.get("TESSERACT_CMD")
                if cmd:
                    pytesseract.pytesseract.tesseract_cmd = cmd
            except Exception:
                pass
            
            txt = pytesseract.image_to_string(img, lang='chi_sim+eng')
            return txt.strip()
        except Exception as e:
            print(f"⚠ pytesseract 解析失败: {e}")
    
    return ""


def parse_txt(file_data: bytes, filename: Optional[str] = None) -> str:
    """解析文本文件"""
    try:
        # 尝试 UTF-8
        try:
            return file_data.decode('utf-8').strip()
        except UnicodeDecodeError:
            # 尝试 GBK
            try:
                return file_data.decode('gbk', errors='ignore').strip()
            except Exception:
                # 尝试其他编码
                return file_data.decode('latin-1', errors='ignore').strip()
    except Exception as e:
        print(f"⚠ 文本文件解析失败: {e}")
        return ""


def parse_document(file_data: bytes, filename: Optional[str] = None, file_type: Optional[str] = None) -> Dict[str, Any]:
    """
    通用文档解析函数
    根据文件扩展名或 MIME 类型自动选择解析器
    
    返回: {
        "content": str,  # 提取的文本内容
        "source_type": str,  # 文档来源类型
        "metadata": dict  # 元数据（如页数、工作表数等）
    }
    """
    if not filename and not file_type:
        return {"content": "", "source_type": "unknown", "metadata": {}}
    
    # 确定文件类型
    ext = ""
    if filename:
        ext = Path(filename).suffix.lower()
    
    metadata = {}
    content = ""
    source_type = "file"
    
    # 根据扩展名选择解析器
    if ext in ('.pdf',) or file_type == 'application/pdf':
        content = parse_pdf(file_data, filename)
        source_type = "pdf"
        metadata["parser"] = "pdfplumber" if PDFPLUMBER_AVAILABLE else "pymupdf"
    
    elif ext in ('.docx',) or file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        content = parse_word(file_data, filename)
        source_type = "word"
        metadata["parser"] = "python-docx"
    
    elif ext in ('.xlsx', '.xls') or file_type in ('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel'):
        content = parse_excel(file_data, filename)
        source_type = "excel"
        metadata["parser"] = "pandas"
    
    elif ext in ('.txt', '.md', '.csv', '.log') or (file_type and 'text/' in file_type):
        content = parse_txt(file_data, filename)
        source_type = "text"
        metadata["parser"] = "direct_read"
    
    elif ext in ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp') or (file_type and file_type.startswith('image/')):
        content = parse_image(file_data, filename)
        source_type = "image"
        metadata["parser"] = "paddleocr" if PADDLEOCR_AVAILABLE else "pytesseract"
    
    else:
        # 未知类型，尝试作为文本处理
        content = parse_txt(file_data, filename)
        source_type = "text"
        metadata["parser"] = "direct_read"
        metadata["warning"] = f"未知文件类型 {ext}，按文本处理"
    
    return {
        "content": content,
        "source_type": source_type,
        "metadata": metadata
    }


def parse_from_database(db_session, table_name: str, columns: Optional[list] = None, limit: int = 1000) -> str:
    """
    从数据库表提取数据
    
    参数:
    - db_session: SQLAlchemy 数据库会话
    - table_name: 表名
    - columns: 要提取的列（None 表示所有列）
    - limit: 最大行数
    """
    try:
        from sqlalchemy import text, inspect
        
        # 构建查询
        if columns:
            cols = ", ".join([f'"{col}"' for col in columns])
            query = f'SELECT {cols} FROM "{table_name}" LIMIT {limit}'
        else:
            query = f'SELECT * FROM "{table_name}" LIMIT {limit}'
        
        result = db_session.execute(text(query))
        rows = result.fetchall()
        
        if not rows:
            return ""
        
        # 获取列名
        column_names = result.keys()
        
        # 转换为文本格式
        text_parts = [f"=== 数据库表: {table_name} ===\n"]
        text_parts.append(" | ".join(column_names))
        text_parts.append("\n" + "-" * 80 + "\n")
        
        for row in rows:
            row_text = " | ".join([str(val) if val is not None else "" for val in row])
            text_parts.append(row_text + "\n")
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"⚠ 数据库提取失败: {e}")
        return ""
